import streamlit as st
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
import PyPDF2
import google.generativeai as genai
from streamlit_option_menu import option_menu
import io
from docx import Document

# ==========================================
# PAGE CONFIGURATION & PREMIUM CSS
# ==========================================
st.set_page_config(page_title="AI Placement & ATS System", layout="wide", page_icon="🎓")

st.markdown("""
<style>
    /* 1. Premium SaaS Background */
    .stApp {
        background-color: #f8fafc;
        background-image: radial-gradient(at 40% 20%, hsla(210,100%,95%,1) 0px, transparent 50%),
                          radial-gradient(at 80% 0%, hsla(189,100%,96%,1) 0px, transparent 50%),
                          radial-gradient(at 0% 50%, hsla(210,100%,96%,1) 0px, transparent 50%);
        background-attachment: fixed;
    }
    
    @media (prefers-color-scheme: dark) {
        .stApp {
            background-color: #0f172a;
            background-image: radial-gradient(at 40% 20%, hsla(220,50%,15%,1) 0px, transparent 50%),
                              radial-gradient(at 80% 0%, hsla(200,60%,15%,1) 0px, transparent 50%);
        }
    }

    /* 2. Page Content Animation */
    @keyframes slideUpFade {
        0% { opacity: 0; transform: translateY(40px); }
        100% { opacity: 1; transform: translateY(0); }
    }
    .stMarkdown, div[data-testid="metric-container"], .stImage {
        animation: slideUpFade 0.8s cubic-bezier(0.16, 1, 0.3, 1) forwards;
    }
    .main .block-container {
        padding-top: 1rem; 
        padding-bottom: 120px; 
    }

    /* 3. The Perfect Circular Floating Chatbot (HIGH VISIBILITY) */
    div[data-testid="stPopover"] {
        position: fixed;
        bottom: 40px;
        right: 40px;
        z-index: 999999;
    }
    
    div[data-testid="stPopover"] > button {
        background: linear-gradient(135deg, #4f46e5 0%, #1e1b4b 100%) !important;
        color: white !important;
        border-radius: 50% !important; 
        width: 75px !important;
        height: 75px !important;
        padding: 0 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        box-shadow: 0 12px 25px rgba(49, 46, 129, 0.5), inset 0 2px 4px rgba(255,255,255,0.3) !important;
        border: 2px solid rgba(255,255,255,0.5) !important; 
        transition: transform 0.3s cubic-bezier(0.4, 0, 0.2, 1), box-shadow 0.3s ease !important;
    }
    
    div[data-testid="stPopover"] > button:hover {
        transform: scale(1.1) translateY(-5px) !important;
        box-shadow: 0 20px 35px rgba(49, 46, 129, 0.7), inset 0 2px 4px rgba(255,255,255,0.5) !important;
    }
    
    div[data-testid="stPopover"] > button p, 
    div[data-testid="stPopover"] > button span {
        font-size: 38px !important;
        margin: 0 !important;
        line-height: 1 !important;
        display: flex !important;
        align-items: center !important;
        justify-content: center !important;
        text-shadow: 0 2px 6px rgba(0,0,0,0.5) !important; 
    }

    /* 4. WIDER CHAT WINDOW */
    div[data-testid="stPopoverBody"] {
        width: 450px !important;
        max-width: 90vw !important;
        border-radius: 20px !important;
        box-shadow: 0 15px 40px rgba(0,0,0,0.15) !important;
    }

    /* 5. Glassmorphism Metric Cards */
    div[data-testid="metric-container"] {
        background: rgba(255, 255, 255, 0.65);
        backdrop-filter: blur(16px);
        border: 1px solid rgba(255, 255, 255, 0.4);
        padding: 25px;
        border-radius: 20px;
        box-shadow: 0 10px 30px rgba(0, 0, 0, 0.03);
        transition: transform 0.3s ease;
    }
    div[data-testid="metric-container"]:hover {
        transform: translateY(-5px);
        box-shadow: 0 15px 35px rgba(0, 0, 0, 0.08);
    }

    /* 6. Highly Professional Images */
    img {
        border-radius: 20px;
        box-shadow: 0 15px 40px rgba(0,0,0,0.12);
        margin-bottom: 30px;
        object-fit: cover;
        width: 100%;
        height: 350px; 
        transition: transform 0.5s ease;
    }
    img:hover {
        transform: scale(1.02);
    }
    
    /* 7. Form Button Styling */
    button[kind="primary"] {
        background: linear-gradient(135deg, #1e293b, #0f172a);
        color: white;
        border-radius: 12px;
        font-weight: 600;
        transition: all 0.2s ease;
    }
    button[kind="primary"]:hover {
        transform: translateY(-2px);
        box-shadow: 0 8px 20px rgba(15, 23, 42, 0.3);
    }
</style>
""", unsafe_allow_html=True)


# ==========================================
# ML MODEL SETUP (Cached)
# ==========================================
@st.cache_resource
def train_model():
    np.random.seed(42)
    n_samples = 500
    data = pd.DataFrame({
        'CGPA': np.random.uniform(5.0, 10.0, n_samples),
        'DSA_Score': np.random.randint(10, 100, n_samples),
        'Internships': np.random.randint(0, 4, n_samples),
        'Projects': np.random.randint(1, 6, n_samples),
    })
    score = (data['CGPA']*10 + data['DSA_Score'] + data['Internships']*20 + data['Projects']*10)
    threshold = score.median()
    data['Placed'] = (score > threshold).astype(int)
    X = data.drop('Placed', axis=1)
    y = data['Placed']
    model = RandomForestClassifier(n_estimators=100, random_state=42)
    model.fit(X, y)
    return model

model = train_model()


# ==========================================
# UTILITY FUNCTIONS (LLM & Word Doc Gen)
# ==========================================
def extract_text_from_pdf(uploaded_file):
    pdf_reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def analyze_resume_with_gemini(api_key, resume_text, job_desc):
    genai.configure(api_key=api_key)
    try:
        llm = genai.GenerativeModel('gemini-3.6-flash')
        prompt = f"""
        You are an expert IT Recruiter ATS system.
        Analyze this candidate's resume against the following job description.
        Job Description: {job_desc}
        Resume: {resume_text}
        Provide your output exactly in these 3 sections using Markdown formatting:
        ### 1. Match Percentage
        ### 2. Missing Keywords
        ### 3. Actionable Feedback
        """
        response = llm.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"API Error: {str(e)}"

def generate_tailored_resume(api_key, resume_text, job_desc):
    genai.configure(api_key=api_key)
    try:
        llm = genai.GenerativeModel('gemini-3.6-flash')
        prompt = f"""
        You are an expert Resume Writer and Career Coach.
        Rewrite and format the following resume to perfectly align with the provided job description.
        Highlight relevant skills and experiences. Maintain truthfulness but optimize keywords for ATS.
        Output ONLY the resume text in clean markdown.

        Job Description: {job_desc}
        Original Resume: {resume_text}
        """
        response = llm.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"API Error: {str(e)}"

def create_word_docx(resume_content):
    doc = Document()
    doc.add_heading("Tailored Professional Resume", 0)
    
    # Simple markdown parser for Word Doc formatting
    for line in resume_content.split('\n'):
        if line.strip():
            if line.startswith('###') or line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('-') or line.startswith('*'):
                doc.add_paragraph(line[1:].strip(), style='List Bullet')
            elif line.startswith('**') and line.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(line.replace('**', '').strip()).bold = True
            else:
                doc.add_paragraph(line.strip())
                
    # Save document to memory buffer
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# ==========================================
# PREMIUM NAVIGATION BAR
# ==========================================
col1, col2 = st.columns([3, 1])
with col1:
    st.markdown("<h2 style='font-weight: 800; color: #0f172a;'>🎓 Career Intelligence Platform</h2>", unsafe_allow_html=True)
with col2:
    st.markdown("<div style='text-align: right; color: #64748b; margin-top: 10px;'><strong>Dev:</strong> Vishal Pandey<br><strong>Roll No:</strong> 2400900100155</div>", unsafe_allow_html=True)

st.write("") 

selected_page = option_menu(
    menu_title=None,
    options=["Predictor Engine", "Resume ATS", "Career Roadmap"],
    icons=["graph-up-arrow", "file-earmark-code", "compass"],
    menu_icon="cast",
    default_index=0,
    orientation="horizontal",
    styles={
        "container": {"padding": "8px", "background-color": "rgba(255,255,255,0.8)", "border-radius": "20px", "box-shadow": "0 4px 15px rgba(0,0,0,0.05)", "backdrop-filter": "blur(10px)"},
        "icon": {"color": "#4f46e5", "font-size": "20px"},
        "nav-link": {"font-size": "16px", "font-weight": "600", "text-align": "center", "margin": "0px 5px", "border-radius": "12px", "--hover-color": "#f1f5f9"},
        "nav-link-selected": {"background-color": "#0f172a", "color": "white"},
    }
)

st.write("")
st.write("")

# ==========================================
# PAGE CONTROLLERS
# ==========================================

if selected_page == "Predictor Engine":
    st.image("https://images.unsplash.com/photo-1551288049-bebda4e38f71?auto=format&fit=crop&w=2000&q=80", use_container_width=True)
    st.title("📊 Placement Probability Engine")
    st.markdown("Analyze your academic metrics to receive a precise probability score and a detailed diagnostic report.")
    
    col1, col2 = st.columns([1.2, 1], gap="large")
    
    with col1:
        with st.form("prediction_form"):
            cgpa = st.slider("Current CGPA", 0.0, 10.0, 7.5, 0.1)
            dsa = st.slider("DSA/Coding Score (out of 100)", 0, 100, 60)
            internships = st.number_input("Number of Internships", 0, 5, 1)
            projects = st.number_input("Number of Major Projects", 0, 10, 2)
            predict_btn = st.form_submit_button("Generate Diagnostic Report", use_container_width=True, type="primary")
            
    with col2:
        if predict_btn:
            features = np.array([[cgpa, dsa, internships, projects]])
            prob = model.predict_proba(features)[0][1] * 100
            
            st.metric("Probability Score", f"{prob:.1f}%")
            
            if prob >= 75:
                st.success("🎉 Excellent! Your profile strongly aligns with tier-1 technical requirements.")
            elif prob >= 50:
                st.warning("⚠️ Average Profile. Technical rounds will require intensive preparation.")
            else:
                st.error("🚨 High Risk. Immediate strategic intervention is required.")

    if predict_btn:
        st.divider()
        st.markdown("### 📈 Strategy Report")
        diag_col1, diag_col2 = st.columns(2)
        
        with diag_col1:
            st.markdown("#### 🎯 Execution Focus")
            if dsa < 65:
                st.error("**Critical: DSA Score**\nSolve 2-3 algorithmic problems daily, focusing on Trees and Graphs.")
            else:
                st.success("**Strength: DSA Score**\nYour algorithmic logic is highly competitive.")
                
            if projects < 2:
                st.warning("**Improvement: Practical Work**\nBuild complex, deployed full-stack applications.")
            else:
                st.success("**Strength: Portfolio**\nYour project volume is excellent.")

        with diag_col2:
            st.markdown("#### 🛤️ Alternative Pathways")
            if internships >= 1 and projects >= 2 and dsa < 60:
                st.info("**Data Analytics & Web Dev**\nMaster SQL, Python, and robust backends.")
            else:
                st.info("**Technical Consulting**\nFocus on Aptitude, Communication, and foundational CS concepts.")


elif selected_page == "Resume ATS":
    st.image("https://images.unsplash.com/photo-1498050108023-c5249f4df085?auto=format&fit=crop&w=2000&q=80", use_container_width=True)
    st.title("📄 AI Resume ATS Scanner & Generator")
    st.markdown("Evaluate your resume against specific job descriptions to bypass automated filters, then let AI auto-generate a tailored version.")
    
    # Initialize Session States for ATS Memory
    if "ats_feedback" not in st.session_state:
        st.session_state.ats_feedback = None
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = ""
    if "job_desc" not in st.session_state:
        st.session_state.job_desc = ""

    gemini_key = st.secrets.get("GEMINI_API_KEY", None)
    if not gemini_key:
        st.warning("⚠️ API Key not found! Please configure GEMINI_API_KEY in your Streamlit Cloud Secrets.")
    
    col1, col2 = st.columns([1, 1], gap="large")
    with col1:
        st.session_state.job_desc = st.text_area("Paste Target Job Description", st.session_state.job_desc or "Software Engineer with Python, SQL, and Data Structures expertise.", height=200)
    with col2:
        uploaded_pdf = st.file_uploader("Upload Resume (PDF)", type=["pdf"])
    
    if st.button("Initialize Deep ATS Scan", type="primary", use_container_width=True):
        if not gemini_key:
            st.error("Cannot run scan: Gemini API Key is missing.")
        elif not uploaded_pdf:
            st.error("Please upload a PDF resume.")
        else:
            with st.spinner("🚀 Extracting text, analyzing semantic matches, and querying LLM..."):
                st.session_state.resume_text = extract_text_from_pdf(uploaded_pdf)
                st.session_state.ats_feedback = analyze_resume_with_gemini(gemini_key, st.session_state.resume_text, st.session_state.job_desc)
                
    # Display Results and Offer Generation
    if st.session_state.ats_feedback:
        st.divider()
        st.subheader("🧠 ATS Evaluation Report")
        st.markdown(st.session_state.ats_feedback)
        
        st.divider()
        st.markdown("### 📝 Need a better match?")
        st.info("Would you like our AI to automatically rewrite and format your resume specifically for this job description?")
        
        if st.button("✨ Yes, Auto-Generate Tailored Resume", use_container_width=True):
            with st.spinner("🤖 AI is rewriting your resume for maximum ATS compatibility..."):
                # 1. Ask LLM to rewrite
                tailored_markdown = generate_tailored_resume(gemini_key, st.session_state.resume_text, st.session_state.job_desc)
                
                # 2. Convert Markdown to .docx
                docx_file = create_word_docx(tailored_markdown)
                
                st.success("✅ Tailored Resume Generated Successfully!")
                
                # 3. Provide Download Button
                st.download_button(
                    label="📥 Download Tailored Resume (.docx)",
                    data=docx_file,
                    file_name="AI_Tailored_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )


elif selected_page == "Career Roadmap":
    st.image("https://images.unsplash.com/photo-1451187580459-43490279c0fa?auto=format&fit=crop&w=2000&q=80", use_container_width=True)
    st.title("🗺️ Comprehensive Career Roadmap")
    st.markdown("A deep-dive, 6-month execution plan engineered for your target industry.")
    
    target_tier = st.selectbox(
        "Select Target Placement Tier:",
        ["Product-Based Companies (SDE)", "Data Analytics & AI Roles", "Service-Based Companies"]
    )
    
    tab1, tab2, tab3 = st.tabs(["Phase 1 (Months 1-2)", "Phase 2 (Months 3-4)", "Phase 3 (Months 5-6)"])
    
    if target_tier == "Product-Based Companies (SDE)":
        with tab1:
            st.write("- **Language Mastery:** Master OOPs (Inheritance, Polymorphism).")
            st.write("- **Basic DSA:** Arrays, Strings, Pointers, Recursion.")
        with tab2:
            st.write("- **Heavy DSA:** Deep dive into Binary Trees, Graphs, DP.")
            st.write("- **Backend Frameworks:** Build scalable APIs (ASP.NET MVC or Node.js).")
        with tab3:
            st.write("- **Project Deployment:** Push 2 major projects to GitHub.")
            st.write("- **CS Fundamentals:** Revise OS, Computer Networks, and DBMS.")
            
    elif target_tier == "Data Analytics & AI Roles":
        with tab1:
            st.write("- **Python Fundamentals:** Deep dive into Python list comprehensions.")
            st.write("- **Data Libraries:** Fluency in NumPy and Pandas.")
        with tab2:
            st.write("- **Visuals:** Matplotlib and Plotly for charts.")
            st.write("- **ML Algorithms:** Regression, Decision Trees, Time-Series forecasting.")
        with tab3:
            st.write("- **End-to-End Project:** Build a multi-page dashboard application.")
            st.write("- **Business Analytics:** Practice guesstimates and case studies.")

    elif target_tier == "Service-Based Companies":
        with tab1:
            st.write("- **Quantitative Aptitude:** Practice percentages, speed-distance.")
            st.write("- **Logical Reasoning:** Syllogisms, pattern recognition.")
        with tab2:
            st.write("- **Subject Revision:** Operating Systems, Digital Electronics.")
            st.write("- **Web Basics:** DOM, HTML, CSS, basic JavaScript.")
        with tab3:
            st.write("- **Communication:** Practice spoken English and articulation.")
            st.write("- **Resume Polish:** Ensure zero grammatical errors.")


# ==========================================
# THE FLOATING RIGHT-SIDE AI CHATBOT (CIRCULAR)
# ==========================================
if "chat_messages" not in st.session_state:
    st.session_state.chat_messages = []

with st.popover("🤖"):
    st.markdown("### 🤖 Placement Mentor")
    
    with st.form("chat_form", clear_on_submit=True):
        user_input = st.text_input("Type your question here...")
        send_btn = st.form_submit_button("Send Question", use_container_width=True)
        
    faq_col1, faq_col2 = st.columns(2)
    if faq_col1.button("Improve DSA?", use_container_width=True, key="faq1"):
        st.session_state.chat_messages.append({"role": "user", "content": "How can I improve my DSA scores?"})
        st.rerun()
    if faq_col2.button("Best projects?", use_container_width=True, key="faq2"):
        st.session_state.chat_messages.append({"role": "user", "content": "What are the best software projects?"})
        st.rerun()
    
    st.divider()
    
    chat_container = st.container(height=380)
    with chat_container:
        for msg in st.session_state.chat_messages:
            st.chat_message(msg["role"]).write(msg["content"])
            
    if (send_btn and user_input) or (len(st.session_state.chat_messages) > 0 and st.session_state.chat_messages[-1]["role"] == "user" and send_btn == False):
        if send_btn and user_input:
            st.session_state.chat_messages.append({"role": "user", "content": user_input})
            st.rerun()
            
        try:
            gemini_key = st.secrets.get("GEMINI_API_KEY", "")
            if not gemini_key:
                with chat_container:
                    st.error("API Key missing.")
            else:
                with chat_container:
                    with st.spinner("Mentor is typing..."):
                        genai.configure(api_key=gemini_key)
                        llm_chat = genai.GenerativeModel('gemini-3.6-flash')
                        last_user_msg = st.session_state.chat_messages[-1]["content"]
                        system_prompt = f"You are a helpful college placement mentor. Answer concisely. Student asks: {last_user_msg}"
                        response = llm_chat.generate_content(system_prompt)
                        st.session_state.chat_messages.append({"role": "assistant", "content": response.text})
                        st.rerun()
                        
        except Exception as e:
            with chat_container:
                st.error("Connection Error. Check API Key.")
